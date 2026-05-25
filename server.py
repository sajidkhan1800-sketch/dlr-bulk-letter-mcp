"""
BulkLetterToApplicant MCP Server
Sends personalised FDA DLR labeling decision letters to applicants via Gmail.
"""

import io
import csv
import json
import smtplib
import tempfile
import os
from datetime import datetime
from email import encoders
from email.mime.base import MIMEBase
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from pathlib import Path
from typing import List, Optional

from docx import Document as DocxDocument
from mcp.server.fastmcp import FastMCP
from pydantic import BaseModel, Field, ConfigDict

# ── constants ─────────────────────────────────────────────────────────────────
BASE_DIR        = Path(__file__).parent
DEFAULT_CSV     = BASE_DIR / "anda_labeling_data.csv"
TEMPLATE_PATH   = BASE_DIR / "DLR_Labeling_Decision_Letter_Template.docx"

SENDER_EMAIL    = "sajidkhan1800@gmail.com"
CC_EMAIL        = "sajidkhan1800@gmail.com"
GMAIL_APP_PASS  = "xnuy lbxn ctnf qrbx"   # baked-in App Password

SMTP_HOST = "smtp.gmail.com"
SMTP_PORT = 587

# Runtime state — CSV rows currently loaded
_loaded_rows: List[dict] = []

# ── MCP server ────────────────────────────────────────────────────────────────
mcp = FastMCP("bulk_letter_mcp")


# ── helpers ───────────────────────────────────────────────────────────────────

def _parse_csv_text(text: str) -> List[dict]:
    """Parse CSV text into list of dicts, stripping whitespace from keys."""
    reader = csv.DictReader(io.StringIO(text.strip()))
    rows = []
    for row in reader:
        rows.append({k.strip(): v.strip() for k, v in row.items()})
    return rows


def _load_default_csv() -> List[dict]:
    """Load the bundled CSV file."""
    with open(DEFAULT_CSV, newline="", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        return [{k.strip(): v.strip() for k, v in row.items()} for row in reader]


def _fill_template(row: dict) -> bytes:
    """Replace [PLACEHOLDER] tokens in the docx template with row values."""
    replacements = {
        "[MM/DD/YYYY]":            datetime.today().strftime("%m/%d/%Y"),
        "[ANDA No.]":              row.get("ANDA No", ""),
        "[Drug Name]":             row.get("Drug Name", ""),
        "[Applicant Name]":        row.get("Applicant Name", ""),
        "[Applicant POC]":         row.get("Applicant POC", ""),
        "[Applicant Address]":     row.get("Applicant Address", ""),
        "[Applicant Email]":       row.get("Applicant Email", ""),
        "[DLR POC]":               row.get("DLR POC", ""),
        "[FDA Labeling Decision]": row.get("FDA Labeling Decision", ""),
        "[DLR POC Email]":         "dlr-reviewer@fda.hhs.gov",
    }
    doc = DocxDocument(TEMPLATE_PATH)
    for para in doc.paragraphs:
        for run in para.runs:
            for token, value in replacements.items():
                if token in run.text:
                    run.text = run.text.replace(token, value)
    for table in doc.tables:
        for trow in table.rows:
            for cell in trow.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        for token, value in replacements.items():
                            if token in run.text:
                                run.text = run.text.replace(token, value)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _send_one_email(row: dict) -> dict:
    """Send a single personalised letter. Returns status dict."""
    applicant = row.get("Applicant Name", "Applicant")
    drug      = row.get("Drug Name", "")
    anda      = row.get("ANDA No", "")
    decision  = row.get("FDA Labeling Decision", "")
    to_email  = row.get("Applicant Email", "")

    try:
        docx_bytes = _fill_template(row)
        safe_drug  = "".join(c if c.isalnum() or c in " _-" else "_" for c in drug)
        filename   = f"DLR_Decision_{anda}_{safe_drug}.docx"

        msg = MIMEMultipart()
        msg["From"]    = f"Sajid Khan <{SENDER_EMAIL}>"
        msg["To"]      = to_email
        msg["CC"]      = CC_EMAIL
        msg["Subject"] = f"FDA DLR Labeling Review Decision — {drug} ({anda})"

        body = (
            f"Dear {applicant},\n\n"
            f"Please find attached the FDA Division of Labeling Review (DLR) decision letter "
            f"for {drug} (ANDA {anda}).\n\n"
            f"Labeling Decision: {decision}\n\n"
            f"Please refer to the attached letter for full details and any required next steps.\n\n"
            f"This message was sent on behalf of the FDA Center for Drug Evaluation and Research, "
            f"Office of Generic Drugs, Division of Labeling Review.\n"
        )
        msg.attach(MIMEText(body, "plain"))

        part = MIMEBase(
            "application",
            "vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        part.set_payload(docx_bytes)
        encoders.encode_base64(part)
        part.add_header("Content-Disposition", f'attachment; filename="{filename}"')
        msg.attach(part)

        recipients = list({to_email, CC_EMAIL})
        with smtplib.SMTP(SMTP_HOST, SMTP_PORT) as server:
            server.starttls()
            server.login(SENDER_EMAIL, GMAIL_APP_PASS)
            server.sendmail(SENDER_EMAIL, recipients, msg.as_string())

        return {"anda": anda, "applicant": applicant, "drug": drug,
                "decision": decision, "to": to_email, "cc": CC_EMAIL,
                "status": "sent", "error": None}

    except Exception as e:
        return {"anda": anda, "applicant": applicant, "drug": drug,
                "decision": decision, "to": to_email, "cc": CC_EMAIL,
                "status": "failed", "error": str(e)}


# ── MCP tools ─────────────────────────────────────────────────────────────────

class LoadCsvInput(BaseModel):
    model_config = ConfigDict(str_strip_whitespace=True, extra="forbid")
    csv_text: Optional[str] = Field(
        default=None,
        description=(
            "Raw CSV text to load. If omitted, the bundled anda_labeling_data.csv is used. "
            "Must include headers: ANDA No, Drug Name, Applicant Name, Applicant POC, "
            "Applicant Email, Applicant Address, DLR POC, FDA Labeling Decision."
        )
    )


@mcp.tool(
    name="dlr_load_csv",
    annotations={
        "title": "Load Applicant CSV",
        "readOnlyHint": False,
        "destructiveHint": False,
        "idempotentHint": True,
        "openWorldHint": False,
    }
)
async def dlr_load_csv(params: LoadCsvInput) -> str:
    """Load applicant data from a CSV into the server's working memory.

    If csv_text is provided, parses it directly. Otherwise loads the bundled default CSV.
    Call this before dlr_preview_applicants or dlr_send_letters.

    Args:
        params (LoadCsvInput):
            - csv_text (Optional[str]): Raw CSV content as a string. Omit to use default.

    Returns:
        str: JSON summary with row count and column names.
    """
    global _loaded_rows
    try:
        if params.csv_text:
            _loaded_rows = _parse_csv_text(params.csv_text)
            source = "provided CSV text"
        else:
            _loaded_rows = _load_default_csv()
            source = "bundled anda_labeling_data.csv"

        required = {"ANDA No", "Drug Name", "Applicant Name", "Applicant POC",
                    "Applicant Email", "Applicant Address", "DLR POC", "FDA Labeling Decision"}
        if _loaded_rows:
            missing = required - set(_loaded_rows[0].keys())
            if missing:
                _loaded_rows = []
                return json.dumps({
                    "success": False,
                    "error": f"Missing required columns: {', '.join(sorted(missing))}"
                }, indent=2)

        return json.dumps({
            "success": True,
            "source": source,
            "rows_loaded": len(_loaded_rows),
            "columns": list(_loaded_rows[0].keys()) if _loaded_rows else [],
            "message": f"Successfully loaded {len(_loaded_rows)} applicant rows from {source}."
        }, indent=2)

    except Exception as e:
        return json.dumps({"success": False, "error": str(e)}, indent=2)


@mcp.tool(
    name="dlr_preview_applicants",
    annotations={
        "title": "Preview Loaded Applicants",
        "readOnlyHint": True,
        "destructiveHint": False,
        "idempotentHint": True,
        "openWorldHint": False,
    }
)
async def dlr_preview_applicants() -> str:
    """List all applicants currently loaded in memory.

    Shows ANDA No, Drug Name, Applicant Name, Applicant POC, FDA Labeling Decision,
    and Applicant Email for each row.

    Returns:
        str: JSON list of applicant records with key fields.
    """
    if not _loaded_rows:
        return json.dumps({
            "success": False,
            "error": "No CSV loaded. Call dlr_load_csv first."
        }, indent=2)

    preview = []
    for row in _loaded_rows:
        preview.append({
            "ANDA No":              row.get("ANDA No", ""),
            "Drug Name":            row.get("Drug Name", ""),
            "Applicant Name":       row.get("Applicant Name", ""),
            "Applicant POC":        row.get("Applicant POC", ""),
            "FDA Labeling Decision":row.get("FDA Labeling Decision", ""),
            "Applicant Email":      row.get("Applicant Email", ""),
            "DLR POC":              row.get("DLR POC", ""),
        })

    return json.dumps({
        "success": True,
        "total": len(preview),
        "applicants": preview
    }, indent=2)


class SendLettersInput(BaseModel):
    model_config = ConfigDict(str_strip_whitespace=True, extra="forbid")
    filter_decision: Optional[str] = Field(
        default=None,
        description=(
            "Optional filter — send only to applicants matching this FDA Labeling Decision. "
            "Valid values: 'Acceptable', 'Acceptable-Minor Deficiencies', "
            "'Acceptable-Major Deficiencies', 'Not Acceptable'. "
            "Omit to send to all loaded applicants."
        )
    )
    filter_anda: Optional[str] = Field(
        default=None,
        description="Optional filter — send only to the applicant with this ANDA number (e.g. 'ANDA204517')."
    )


@mcp.tool(
    name="dlr_send_letters",
    annotations={
        "title": "Send DLR Decision Letters",
        "readOnlyHint": False,
        "destructiveHint": False,
        "idempotentHint": False,
        "openWorldHint": True,
    }
)
async def dlr_send_letters(params: SendLettersInput) -> str:
    """Send personalised FDA DLR labeling decision letters to all (or filtered) loaded applicants.

    Each applicant receives an email with their filled .docx letter attached.
    Emails are sent FROM sajidkhan1800@gmail.com and CC'd TO sajidkhan1800@gmail.com.
    Call dlr_load_csv before this tool.

    Args:
        params (SendLettersInput):
            - filter_decision (Optional[str]): Only send to applicants with this decision value.
            - filter_anda (Optional[str]): Only send to the applicant with this ANDA number.

    Returns:
        str: JSON summary with per-applicant send status, counts of sent/failed.
    """
    if not _loaded_rows:
        return json.dumps({
            "success": False,
            "error": "No CSV loaded. Call dlr_load_csv first."
        }, indent=2)

    rows = _loaded_rows[:]

    # Apply filters
    if params.filter_anda:
        rows = [r for r in rows if r.get("ANDA No", "").strip() == params.filter_anda.strip()]
        if not rows:
            return json.dumps({
                "success": False,
                "error": f"No applicant found with ANDA No '{params.filter_anda}'."
            }, indent=2)

    if params.filter_decision:
        rows = [r for r in rows if r.get("FDA Labeling Decision", "").strip() == params.filter_decision.strip()]
        if not rows:
            return json.dumps({
                "success": False,
                "error": f"No applicants found with decision '{params.filter_decision}'."
            }, indent=2)

    results = [_send_one_email(row) for row in rows]
    sent    = sum(1 for r in results if r["status"] == "sent")
    failed  = sum(1 for r in results if r["status"] == "failed")

    return json.dumps({
        "success": True,
        "total":   len(results),
        "sent":    sent,
        "failed":  failed,
        "from":    SENDER_EMAIL,
        "cc":      CC_EMAIL,
        "results": results
    }, indent=2)


@mcp.tool(
    name="dlr_status",
    annotations={
        "title": "Check Server Status",
        "readOnlyHint": True,
        "destructiveHint": False,
        "idempotentHint": True,
        "openWorldHint": False,
    }
)
async def dlr_status() -> str:
    """Check the current state of the BulkLetterToApplicant MCP server.

    Returns template availability, loaded CSV row count, sender config, and available tools.

    Returns:
        str: JSON object with server status details.
    """
    return json.dumps({
        "server":            "bulk_letter_mcp",
        "template_ready":    TEMPLATE_PATH.exists(),
        "template_file":     TEMPLATE_PATH.name,
        "default_csv_ready": DEFAULT_CSV.exists(),
        "default_csv_file":  DEFAULT_CSV.name,
        "rows_loaded":       len(_loaded_rows),
        "sender_email":      SENDER_EMAIL,
        "cc_email":          CC_EMAIL,
        "smtp_host":         SMTP_HOST,
        "available_tools": [
            "dlr_status          — check server state",
            "dlr_load_csv        — load applicant CSV (default or custom text)",
            "dlr_preview_applicants — list all loaded applicants",
            "dlr_send_letters    — send letters (optional filters: filter_anda, filter_decision)",
        ]
    }, indent=2)


# ── entry point ───────────────────────────────────────────────────────────────
if __name__ == "__main__":
    mcp.run(transport="stdio")
